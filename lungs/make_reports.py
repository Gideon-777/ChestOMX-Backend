import webpage2html
import base64
import msgpack
from mako.template import Template
import re
import json
import os
import argparse
import k3d_pro as k3d
from vtk.util import numpy_support
import pyacvd
import pyvista
import vtk
import colorsys
import SimpleITK as sitk
# import pdfkit
import numpy as np
import yaml
from k3d_pro.headless import k3d_remote#, get_headless_driver
import subprocess
import copy
from docxtpl import DocxTemplate
import docx
import datetime

def maybe_mkdir(path):
    if not os.path.exists(path):
        os.mkdir(path)

def get_headless_driver(no_headless=False):
    from selenium.webdriver.chrome.options import Options
    from selenium import webdriver
    from webdriver_manager.chrome import ChromeDriverManager

    options = Options()

    if not no_headless:
        options.add_argument("--headless")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")

    return webdriver.Chrome(ChromeDriverManager().install(), options=options)

## widgets

cm = []
for i in range(254):
    c = k3d.nice_colors[i % 19]
    cm += [(c >> (8 * (2-i))) & 255 for i in range(3)]

def save_dicom(prefix, dicoms, folder_name, scan=None, mask=None, series_number=1):
    writer = sitk.ImageFileWriter()
    writer.KeepOriginalImageUIDOn()
    writer.UseCompressionOn()

    modification_time = time.strftime("%H%M%S")
    modification_date = time.strftime("%Y%m%d")

    if scan is not None and mask is not None:
        d = sitk.GetArrayFromImage(scan)

    for idx, image in enumerate(dicoms):
        if scan is not None and mask is not None:
            image_255 = sitk.Cast(sitk.IntensityWindowing(image,
                                                          windowMinimum=np.percentile(d, 0),
                                                          windowMaximum=np.percentile(d, 95),
                                                          outputMinimum=0.0,
                                                          outputMaximum=255.0), sitk.sitkUInt8)

            m = sitk.GetImageFromArray(np.expand_dims(mask[idx, :, :], axis=0))
            m.CopyInformation(image)

            final = sitk.LabelOverlay(image=image_255,
                                      labelImage=m,
                                      opacity=0.5, backgroundValue=0.0,
                                      colormap=cm)
        else:
            final = scan[idx]

        series_tag_values = [
            ("0008|0031", modification_time),  # Series Time
            ("0008|0012", modification_date),  # Instance Creation Date
            ("0008|0013", modification_time),  # Instance Creation Time
            ("0008|0021", modification_date),  # Series Date
            ("0008|103E", "[REPORT] " + prefix),  # Series Description
            ("0020|4000", "[REPORT]"),            # Image Comments
            ("0020|000e", "1.2.826.0.1.3680043.2.1125." + modification_date + ".1" + modification_time),
            ("0020|0013", str(idx)),
            ("0020|0011", str(series_number))
        ]

        tags_to_copy = [#"0010|0010",  # Patient Name
                        "0010|0020",  # Patient ID
                        "0010|0030",  # Patient Birth Date
                        "0020|000d",  # Study Instance UID, for machine consumption
                        "0020|0010",  # Study ID, for human consumption
                        "0008|0020",  # Study Date
                        "0008|0030",  # Study Time
                        "0008|0050",  # Accession Number
                        "0008|0060",  # Modality,
                        "0020|0032",
                        "0020|1041",
                        "0008|1030"  # study description
                        ]

        for tag in tags_to_copy:
            if tag in image.GetMetaDataKeys():
                try:
                    final.SetMetaData(tag, image.GetMetaData(tag))
                except:
                    final.SetMetaData(tag, "UNREADABLE")

        list(map(lambda tag_value: final.SetMetaData(tag_value[0], tag_value[1]), series_tag_values))

        if series_number > 1:
            final.SetMetaData("0008|0008", "DERIVED\\SECONDARY")

        if not os.path.exists(folder_name):
            os.makedirs(folder_name)

        writer.SetFileName(folder_name + "/%s_%06d.dcm" % (prefix, idx))
        writer.Execute(final)

def get_mesh_plot(poly):
    plot = k3d.plot(height=400, lighting=1.75)
    plot.menu_visibility = False
    plot.axes_helper = False

    for mesh in poly.values():
        p = k3d.mesh(vertices=mesh.vertices, indices=mesh.indices, color=mesh.color, name=mesh.name,
                     opacity=mesh.opacity, flat_shading=False)
        p.hover_callback = mesh.hover_callback

        p.model_matrix = k3d.transform(translation=np.random.random(3) * 2).model_matrix
        plot += p

    plot.grid_visible = False
    plot.grid_auto_fit = True
    plot.depth_peels = 12
    plot.camera_auto_fit = False
    plot.camera = plot.get_auto_camera(0.8, -35, 70)

    return plot


def get_3d_plot(scan, masks, poly, size_reduce=[1, 1, 1], poly3d=None):
    plot = k3d.plot(height=400, background_color=0, lighting=1.75)
    plot.menu_visibility = False
    plot.axes_helper = False

    data = sitk.GetArrayFromImage(scan).astype(np.float32).copy()
    data = data[::size_reduce[0], ::size_reduce[1], ::size_reduce[2]]

    volume_slice = k3d.volume_slice(
        data,
        color_range=[-1024, np.percentile(data, 98)],
        color_map=k3d.matplotlib_color_maps.Binary_r,
        bounds=get_bounds(scan),
        slice_z=data.shape[0] // 2,
        slice_y=data.shape[1] // 2,
        slice_x=data.shape[2] // 2,
        mask=masks[::size_reduce[0], ::size_reduce[1], ::size_reduce[2]],
        active_masks=[],
        mask_opacity=0.5,
        name='scan',
        color_map_masks=[k3d.nice_colors[i % 19] for i in range(256)]
    )

    for mesh in poly.values():
        p = k3d.mesh(vertices=mesh.vertices, indices=mesh.indices, color=mesh.color, name=mesh.name)
        p.visible = False
        plot += p

    plot += volume_slice

    plot.slice_viewer_mask_object_ids = [o.id for o in plot.objects if o.type == 'Mesh']
    plot.grid_visible = False
    plot.slice_viewer_object_id = volume_slice.id
    plot.grid_auto_fit = True
    plot.camera_auto_fit = False

    # plot.hidden_object_ids = plot.object_ids

    if poly3d is None:
        poly3d = poly

    for mesh in poly3d.values():
        p = k3d.mesh(vertices=mesh.vertices, indices=mesh.indices, color=mesh.color,
                     opacity=mesh.opacity, name='mesh_' + mesh.name, flat_shading=False)
        plot += p

    # v = np.vstack([p.vertices for p in poly3d.values()])
    # center = np.mean(v, axis=0).tolist()
    # size = (np.max(v, axis=0) - np.min(v, axis=0)).tolist()
    # plot.camera = [
    #     center[0] + size[0] * 0.85, center[1] - size[1] * 1.35, center[2] + size[2] * 0.3,
    #     center[0], center[1], center[2],
    #     0, 0, 1
    # ]

    return plot


def get_2d_plot(scan, masks, poly, size_reduce=[1, 1, 1]):
    plot = k3d.plot(height=400, background_color=0, camera_mode='slice_viewer')

    plot.menu_visibility = False
    plot.axes_helper = False

    data = sitk.GetArrayFromImage(scan).astype(np.float32).copy()
    data = data[::size_reduce[0], ::size_reduce[1], ::size_reduce[2]]

    volume_slice = k3d.volume_slice(
        data,
        color_range=[-1024, np.percentile(data, 98)],
        color_map=k3d.matplotlib_color_maps.Binary_r,
        color_map_masks=[k3d.nice_colors[i % 19] for i in range(256)],
        bounds=get_bounds(scan), slice_z=data.shape[0] // 2,
        mask=masks[::size_reduce[0], ::size_reduce[1], ::size_reduce[2]],
        active_masks=np.unique(masks).tolist(),
        mask_opacity=0.5,
        name='scan'
    )

    for p in poly.values():
        plot += p

    plot += volume_slice

    plot.slice_viewer_mask_object_ids = [o.id for o in plot.objects if o.type == 'Mesh']
    plot.grid_visible = False
    plot.slice_viewer_object_id = volume_slice.id
    plot.grid_auto_fit = True
    plot.camera_auto_fit = False

    return plot

def get_screenshots(data, name="data", port=8080):
    ret = {}

    plot = k3d.plot(grid_visible=False, background_color=0,
                    camera_mode='slice_viewer', axes_helper=0,
                    screenshot_scale=3)

    plot.minimum_fps = -1
    headless = k3d_remote(plot, get_headless_driver(), 1024, 1024, port=port)
    headless.sync(hold_until_refreshed=True)

    for o in data.objects:
        if o['type'] == 'VolumeSlice':
            # obj = o.clone()
            obj = copy.deepcopy(o)
            obj.compression_level = 0
            obj.color_range = [-1300, obj.color_range[1]]
            obj.active_masks = np.unique(obj.mask)[1:]
            plot += obj

    for axis in ['x', 'y', 'z']:
        plot.slice_viewer_direction = axis

        if len(plot.objects) > 0:
            plot.objects[0].slice_x = -1
            plot.objects[0].slice_y = -1
            plot.objects[0].slice_z = -1

            if axis == 'x':
                plot.objects[0].slice_x = plot.objects[0].volume.shape[2] // 4  # chest specific

            if axis == 'y':
                plot.objects[0].slice_y = plot.objects[0].volume.shape[1] // 2

            if axis == 'z':
                plot.objects[0].slice_z = plot.objects[0].volume.shape[0] // 2

        headless.sync(hold_until_refreshed=True)
        ret[name + "_" + axis] = headless.get_screenshot()

    while (len(plot.objects) > 0):
        plot -= plot.objects[-1]

    plot.camera_mode = 'trackball'
    headless.sync(hold_until_refreshed=True)

    for o in data.objects:
        if o.name.startswith("mesh_") or (name == 'Summary' and o['type'] == 'Mesh'):
            # plot += o.clone()
            plot += copy.deepcopy(o)

    if len(plot.objects) > 0:
        plot.camera_auto_fit = False
        plot.colorbar_object_id = 0
        plot.camera = plot.get_auto_camera(0.8, -35, 85)
        headless.sync(hold_until_refreshed=True)
        ret[name + '_3d'] = headless.get_screenshot()

    headless.close()

    return ret

def rearange_masks(masks, scan):
    masks = sitk.GetArrayFromImage(masks)
    new_masks = np.zeros_like(masks)

    for i, v in enumerate(np.unique(masks)):
        new_masks[masks == v] = i

    masks = sitk.GetImageFromArray(new_masks)
    masks.CopyInformation(scan)

    return masks


def load_unify_image(path=None, image=None):
    
    image = sitk.ReadImage(path)

    image.SetDirection(np.round(np.array(image.GetDirection())))

    direction = np.array(image.GetDirection()).reshape(len(image.GetSize()), -1)
    ind = np.argmax(np.abs(direction), axis=1)

    new_size = np.array(image.GetSize())[ind]
    new_spacing = np.array(image.GetSpacing())[ind]
    new_extent = (new_size - np.ones(3)) * new_spacing
    new_origin = np.array(image.GetOrigin()) - new_extent * (np.diag(direction[:, ind]) < 0)

    resample = sitk.ResampleImageFilter()

    resample.SetOutputSpacing(new_spacing.tolist())
    resample.SetSize(new_size.tolist())
    resample.SetOutputDirection(np.eye(3).flatten())
    resample.SetOutputOrigin(new_origin.tolist())

    resample.SetTransform(sitk.Transform())
    resample.SetDefaultPixelValue(image.GetPixelIDValue())
    resample.SetInterpolator(sitk.sitkNearestNeighbor)

    # pad to divide by 4
    im = resample.Execute(image)
    pad = [x % 4 for x in im.GetSize()]
    im = sitk.ConstantPad(im, [0, 0, 0], pad)
    im.SetOrigin((0, 0, 0))

    return im


def scale_lightness(rgb, scale_l):
    r, g, b = (rgb // 256 // 256 % 256 / 255.0, rgb // 256 % 256 / 255.0, rgb % 256 / 255.0)
    h, l, s = colorsys.rgb_to_hls(r, g, b)
    r, g, b = colorsys.hls_to_rgb(h, min(1, l * scale_l), s=s)

    return (int(r * 255.0) << 16) + (int(g * 255.0) << 8) + int(b * 255.0)


def get_id_by_name(name, d, key='Structure'):
    for k, i in d.items():
        if i[key] == name:
            return int(k)


def get_bounds(img):
    origin = img.GetOrigin()
    size = np.array(img.GetSpacing()) * np.array(img.GetSize())

    return np.array([origin[0], origin[0] + size[0],
                     origin[1], origin[1] + size[1],
                     origin[2], origin[2] + size[2]])


def get_volume(img):
    data = sitk.GetArrayFromImage(img).astype(np.float32)

    return k3d.volume(data, bounds=get_bounds(img),
                      samples=512, color_map=k3d.paraview_color_maps.Jet, shadow='dynamic',
                      shadow_delay=250)


def resample(image, size):
    new_spacing = np.array(image.GetSpacing()) * (np.array(image.GetSize()) / np.array(size))

    return sitk.Resample(image, size, sitk.Transform(), sitk.sitkLinear, image.GetOrigin(),
                         new_spacing, image.GetDirection(), 0.0, image.GetPixelIDValue())


def contour(data, bounds, values, quantize_points_factor=0.0, clustering_factor=6):
    vti = vtk.vtkImageData()
    vti.SetOrigin(bounds[::2])
    vti.SetDimensions(np.array(data.shape[::-1]))
    vti.SetSpacing((bounds[1::2] - bounds[::2]) / np.array(data.shape[::-1]))

    arr = numpy_support.numpy_to_vtk(num_array=data.ravel(), deep=True, array_type=vtk.VTK_FLOAT)
    arr.SetName('d')
    vti.GetPointData().AddArray(arr)
    vti.GetPointData().SetActiveScalars('d')

    contour = vtk.vtkImageMarchingCubes()
    contour.SetInputData(vti)
    contour.SetNumberOfContours(values.shape[0])

    for i, v in enumerate(values):
        contour.SetValue(i, v)
    contour.ComputeScalarsOn()
    contour.Update()

    quantize = vtk.vtkQuantizePolyDataPoints()
    quantize.SetQFactor(quantize_points_factor)

    quantize.SetInputConnection(contour.GetOutputPort())
    quantize.Update()

    if clustering_factor > 0:
        mesh = pyvista.PolyData(quantize.GetOutput())
        mesh.compute_normals(inplace=True)
        clus = pyacvd.Clustering(mesh)
        clus.cluster(mesh.n_points // clustering_factor)

        try:
            poly = clus.create_mesh()
        except:
            poly = quantize.GetOutput()
    else:
        poly = quantize.GetOutput()

    return poly

# add by mulder
def flipMsk12(img):
    return sitk.GetImageFromArray(np.flip(np.flip(sitk.GetArrayFromImage(img) , 1), 2))







## plots


def get_poly(masks, prefix, size_reduce=[1, 1, 1], smooth=18, names=None, ignore=[]):
    poly = {}
    img = sitk.GetArrayFromImage(masks).astype(np.float32)
    bounds = get_bounds(masks)
    volumes = {}

    for i in np.unique(img)[1:]:
        center = np.median(np.dstack(np.where(img == i)), axis=1)[0] // np.array(size_reduce)

    #    if names is None or str(int(i)) not in names.keys():
    #        name = prefix + ' ' + str(int(i))
    #    else:
    #        name = names[str(int(i))]
        j = int(i - 1)
        name = names[j]

        if name in ignore:
            continue

        volumes[str(i)] = {
            'Label': name,
            'name': name.replace(' ', '_'),
            'id': [int(i)],
            'center': center[::-1].astype(np.int32).tolist(),
            'Volume': np.sum(img == i) * np.prod(np.array(masks.GetSpacing())) / (10 ** 3)
        }

    for idx, val in list(enumerate(np.unique(img)[1:])):
        p = contour(img == val, bounds, np.array([0.5]), 0.0, smooth)

        name = prefix + '_' + str(int(val))
        poly[name] = k3d.vtk_poly_data(p, name=name, opacity=1.0, color=k3d.nice_colors[idx + 1])

    return poly, img.astype(np.uint8), volumes


def versiontuple(v):
    return tuple(map(int, (v.split("."))))


assert versiontuple(k3d.__version__) >= versiontuple("2.12.0")

IPF_map = {
    "1": "Normal",
    "2": "GroundGlass",
    "3": "Honeycomb",
    "4": "Reticular",
    "5": "ModerateLAA",
    "6": "MildLAA",
    "7": "SevereLAA"
}

def make_report(
    patient_id,
    unique_id,
    models,
    model_labels,
    ignore,
    niftis_folder,
    uploads_folder,
    outputs_folder,
    templates_folder,
    reports_folder,
    patient_info,
    ctss=None
):
    mask_flip = False
    summary = True
    reduce_x = 1
    reduce_y = 1
    reduce_z = 1
    dicom = False

    variables = {}
    series_number = 2
    plots = []
    exam = f"{patient_id}_{unique_id}"
    filename = exam
    variables['case'] = exam

    variables['dicom'] = dicom
    # variables['lobes3d'] = 'lobes' in models
    # variables['lobes'] = 'lobes' in models
    # variables['lungs'] = 'lungs' in models
    # # variables['lobes_ipf_3d'] = args.lobes_ipf_3d
    # variables['ipf'] = 'ild' in models
    # variables['ggo_consolidation'] = args.ggo_consolidation
    for model in models:
        variables[model] = True
    if 'lobes' in models: variables['lobes3d'] = True
    variables['models'] = models

    # image loading
    image_path = f"{niftis_folder}/images/{patient_id}_{unique_id}_0000.nii.gz"
    scan = load_unify_image(image_path)
    dicoms = None
    size_reduce = [reduce_z, reduce_y, reduce_x]
    print(scan.GetSize())
    scan.SetOrigin((0, 0, 0))
    voxel_volume = np.prod(np.array(scan.GetSpacing())) / (10 ** 3)

    possible_dicoms_folder = f"{uploads_folder}/dcm/{patient_id}_{unique_id}"

    if os.path.isdir(possible_dicoms_folder):
        reader = sitk.ImageSeriesReader()
        reader.MetaDataDictionaryArrayUpdateOn()
        reader.LoadPrivateTagsOn()

        dicom_names = reader.GetGDCMSeriesFileNames(possible_dicoms_folder)
        dicoms = [sitk.ReadImage(f) for f in dicom_names]

    if dicom and dicoms is not None:
        print('dicom')
        reader = dicoms[0]

        # many magic strings - it's a dicom standard - https://dicom.innolitics.com/ciods/rt-plan/patient/00100010
        meta = {
            "patientName": reader.GetMetaData('0010|0010') \
                if '0010|0010' in reader.GetMetaDataKeys() else '',
            "patientId": reader.GetMetaData('0010|0020') \
                if '0010|0020' in reader.GetMetaDataKeys() else '',
            "patientBirthdate": reader.GetMetaData('0010|0030') \
                if '0010|0030' in reader.GetMetaDataKeys() else '',
            "patientSex": reader.GetMetaData('0010|0040') \
                if '0010|0040' in reader.GetMetaDataKeys() else '',
            "patientAge": reader.GetMetaData('0010|1010') \
                if '0010|1010' in reader.GetMetaDataKeys() else '',
            "studyDescription": reader.GetMetaData('0008|1030') \
                if '0008|0030' in reader.GetMetaDataKeys() else '',
            "institutionName": reader.GetMetaData('0008|0080') \
                if '0008|0080' in reader.GetMetaDataKeys() else '',
            "institutionAddress": reader.GetMetaData('0008|0081') \
                if '0008|0081' in reader.GetMetaDataKeys() else '',
            "stationName": reader.GetMetaData('0008|1010') \
                if '0008|1010' in reader.GetMetaDataKeys() else '',
            "manufacturer": reader.GetMetaData('0008|1090') \
                if '0008|1090' in reader.GetMetaDataKeys() else '',
            "sliceThickness": reader.GetMetaData('0018|0050') \
                if '0018|0050' in reader.GetMetaDataKeys() else ''
        }

        variables = {**variables, **meta}

    variables['patientName'] = patient_info['patientName'] if 'patientName' in patient_info else ''
    variables['patientId'] = patient_info['patientId'] if 'patientId' in patient_info else ''
    variables['patientBirthDate'] = patient_info['patientBirthDate'] if 'patientBirthDate' in patient_info else ''
    variables['patientSex'] = patient_info['patientSex'] if 'patientSex' in patient_info else ''
    variables['patientAge'] = patient_info['patientAge'] if 'patientAge' in patient_info else ''
    variables['institutionName'] = patient_info['institutionName'] if 'institutionName' in patient_info else ''
    variables['seriesDesc'] = patient_info['seriesDesc'] if 'seriesDesc' in patient_info else ''
    variables['predictedAt'] = patient_info['predictedAt'] if 'predictedAt' in patient_info else datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    if len(variables['predictedAt']) < 2: 
        variables['predictedAt'] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")


    if ctss is not None:
        variables['ctss'] = True
        variables['ctss_lobes'] = [
            'LUL', 'LLL', 'RUL', 'RML', 'RLL'
        ]
        variables['ctss_inv'] = [
            round(ctss['involvement_' + lobe.lower()], 2) for lobe in variables['ctss_lobes']
        ]
        variables['ctss_ctss'] = [
            ctss['ctss_' + lobe.lower()] for lobe in variables['ctss_lobes']
        ]


    # if 'lungs' in models:
    #     print('lungs')
    #     lungs = load_unify_image(f"{outputs_folder}/lungs/{patient_id}_{unique_id}.nii.gz")
    #     if args.mask_flip:
    #         lungs = flipMsk12(lungs)
    #     lungs.SetOrigin((0, 0, 0))
    #     lungs.SetSpacing(scan.GetSpacing())

    #     poly_lungs_smoothed, masks_lungs, lungs_volumes = get_poly(lungs, 'Lungs', size_reduce, 18)
    #     poly_lungs, masks_lungs, lungs_volumes = get_poly(lungs, 'Lungs', size_reduce, 2)

    #     lungs_plot = get_3d_plot(scan, masks_lungs, poly_lungs, size_reduce, poly_lungs_smoothed)

    #     plots.append((lungs_plot, 'lungs', 'volumeViewer'))

    #     variables['lungs_volumes'] = [lungs_volumes[i] for i in sorted(lungs_volumes.keys())]

    #     if dicoms is not None:
    #         save_dicom('lungs', dicoms, filename, scan, masks_lungs, series_number)
    #         series_number += 1

    # if 'lobes' in models:
    #     print('lobes common')
    #     lobes = load_unify_image(f"{outputs_folder}/lobes/{patient_id}_{unique_id}.nii.gz")
    #     if mask_flip:
    #         lobes = flipMsk12(lobes)
    #     lobes.SetOrigin((0, 0, 0))
    #     lobes.SetSpacing(scan.GetSpacing())
    #     poly_lobes_smoothed, masks_lobe, lobes_volumes = get_poly(lobes, 'Lobes', size_reduce, 18)

    #     if dicoms is not None:
    #         save_dicom('lobes', dicoms, filename, scan, masks_lobe, series_number)
    #         series_number += 1

    #     print('lobes3d')
    #     lobes3d_plot = get_mesh_plot(poly_lobes_smoothed)

    #     plots.append((lobes3d_plot, 'lobes3d', 'meshViewer'))

    #     variables['lobes_volumes'] = [lobes_volumes[i] for i in sorted(lobes_volumes.keys())]

    #     print('lobes')
    #     poly_lobes, masks_lobe, lobes_volumes = get_poly(lobes, 'Lobes', size_reduce, 8)
    #     lobes_plot = get_3d_plot(scan, masks_lobe, poly_lobes, size_reduce, poly_lobes_smoothed)

    #     plots.append((lobes_plot, 'lobes', 'volumeViewer'))

    #     variables['lobes_volumes'] = [lobes_volumes[i] for i in sorted(lobes_volumes.keys())]



    # # if args.ipf and (args.lobes or args.lobes3d):
    # if 'ild' in models and 'lobes' in models:
    #     print('ipf')
    #     IPF = load_unify_image(f"{outputs_folder}/ild/{patient_id}_{unique_id}.nii.gz")
    #     if mask_flip:
    #         IPF = flipMsk12(IPF)
    #     IPF.SetOrigin((0, 0, 0))
    #     IPF.SetSpacing(scan.GetSpacing())

    #     poly_IPF, masks_IPF, IPF_volumes = get_poly(IPF, 'IPF', size_reduce, 8, names=IPF_map)
    #     IPF_plot = get_3d_plot(scan, masks_IPF, poly_IPF, size_reduce)

    #     plots.append((IPF_plot, 'IPF', 'volumeViewer'))

    #     variables['IPF_volumes'] = [IPF_volumes[i] for i in sorted(IPF_volumes.keys())]
    #     variables['IPF'] = variables['IPF_volumes']

    #     if dicoms is not None:
    #         save_dicom('IPF', dicoms, filename, scan, masks_IPF, series_number)
    #         series_number += 1

    # # if args.ipf and (args.lobes or args.lobes3d) and args.summary:
    # if 'ild' in models and 'lobes' in models and 'summary' in models:
    #     print('summary')
    #     variables['summary'] = []

    #     for l in range(1, 6):
    #         row = {}
    #         lobe_volume = np.sum(masks_lobe == l) * voxel_volume

    #         for ipf_id, ipf_name in IPF_map.items():
    #             row[IPF_map[ipf_id]] = np.sum((masks_lobe == l) * (masks_IPF == int(ipf_id))) * voxel_volume

    #         lobe = {
    #             'Label': 'Lobes ' + str(l),
    #             'Volume': lobe_volume,
    #             'volumes': row
    #         }

    #         variables['summary'].append(lobe)
    variables['model_volumes'] = {}
    for model in models: # lungs, lobes, ild...

        output_img = load_unify_image(f"{outputs_folder}/{model}/{patient_id}_{unique_id}.nii.gz")
        if mask_flip:
            output_img = flipMsk12(output_img)
        output_img.SetOrigin((0, 0, 0))
        output_img.SetSpacing(scan.GetSpacing())

        

        poly_output_smoothed, masks, output_volumes = get_poly(output_img, model.capitalize(), size_reduce, 18, names=list(model_labels[model].keys()), ignore=ignore[model] if model in ignore else [])
        poly_output, masks, output_volumes = get_poly(output_img, model.capitalize(), size_reduce, 8, names=list(model_labels[model].keys()), ignore=ignore[model] if model in ignore else [])

        if model == 'lobes':
            lobes3d_plot = get_mesh_plot(poly_output_smoothed)
            plots.append((lobes3d_plot, 'lobes3d', 'meshViewer'))

        output_plot = get_3d_plot(scan, masks, poly_output, size_reduce, poly_output_smoothed)
        plots.append((output_plot,  model, 'volumeViewer'))

        variables[f'{model}_volumes'] = [output_volumes[i] for i in sorted(output_volumes.keys())]
        variables['model_volumes'][model] = [output_volumes[i] for i in sorted(output_volumes.keys())]

        # if dicoms is not None:
        #     save_dicom('lungs', dicoms, filename, scan, masks_lungs, series_number)
        #     series_number += 1

    if summary and ('lobes' in models and 'fibrosis_7labels' in models):

        variables['summary'] = []

        lobes = load_unify_image(f"{outputs_folder}/lobes/{patient_id}_{unique_id}.nii.gz")
        if mask_flip:
            lobes = flipMsk12(lobes)
        lobes.SetOrigin((0, 0, 0))
        lobes.SetSpacing(scan.GetSpacing())
        poly_lobes_smoothed, masks_lobe, lobes_volumes = get_poly(lobes, 'Lobes', size_reduce, 18, names=list(model_labels['lobes'].keys()))

        IPF = load_unify_image(f"{outputs_folder}/fibrosis_7labels/{patient_id}_{unique_id}.nii.gz")
        if mask_flip:
            IPF = flipMsk12(IPF)
        IPF.SetOrigin((0, 0, 0))
        IPF.SetSpacing(scan.GetSpacing())

        poly_IPF, masks_IPF, IPF_volumes = get_poly(IPF, 'IPF', size_reduce, 8, names=list(model_labels['fibrosis_7labels'].keys()))

        for l in range(1, 6):
            row = {}
            lobe_volume = np.sum(masks_lobe == l) * voxel_volume

            for ipf_id, ipf_name in IPF_map.items():
                row[IPF_map[ipf_id]] = np.sum((masks_lobe == l) * (masks_IPF == int(ipf_id))) * voxel_volume

            lobe = {
                'Label': 'Lobes ' + str(l),
                'Volume': lobe_volume,
                'volumes': row
            }

            variables['summary'].append(lobe)



    # if args.lobes_ipf_3d and (args.lobes or args.lobes3d) and args.ipf:
    #     print('lobes_ipf_3d')
    #
    #     poly_IPF, _, _ = get_poly(IPF, 'Infiltration', 2)
    #
    #     for p in poly_IPF.values():
    #         p.hover_callback = True
    #         p.color = 0xff0000
    #
    #     for p in poly_lobes_smoothed.values():
    #         p.opacity = 0.3
    #
    # label = k3d.label("test", name="label", is_html=True, color=0, size=2.0, mode='local')
    # label.visible = False
    #
    #     lobes_infiltration_3d_plot = get_mesh_plot({**poly_lobes_smoothed, **poly_IPF})
    #     lobes_infiltration_3d_plot += label
    #     lobes_infiltration_3d_plot.mode = 'callback'
    #
    #     plots.append((lobes_infiltration_3d_plot, 'lobes_infiltration_3d', 'meshViewer'))

    # if args.ggo_consolidation and (args.lobes or args.lobes3d):
    #     print('ggo_consolidation')
    #     ggo_consolidation, _ = load_unify_image(
    #         './input/ggo_consolidation/' + args.exam + '*.nii.gz')
    #     ggo_consolidation.SetOrigin((0, 0, 0))
    #     ggo_consolidation.SetSpacing(scan.GetSpacing())
    #
    #     poly_ggo_consolidation, masks_ggo_consolidation, ggo_consolidation_volumes = \
    #         get_poly(ggo_consolidation, 'GGO_Consolidation', 2)
    #
    #     ggo_consolidation_plot = get_3d_plot(scan, masks_ggo_consolidation, poly_ggo_consolidation,
    #                                          size_reduce)
    #
    #     plots.append((ggo_consolidation_plot, 'ggo_consolidation', 'volumeViewer'))
    #
    #     variables['ggo'] = {
    #         'LeftGGO': {
    #             'Label': 'Left ggo',
    #             'Volume': np.sum((masks_lobe < 3) * (masks_ggo_consolidation == 2)) * voxel_volume
    #         },
    #         'RightGGO': {
    #             'Label': 'Right ggo',
    #             'Volume': np.sum((masks_lobe >= 3) * (masks_ggo_consolidation == 2)) * voxel_volume
    #         },
    #         'LeftConsolidation': {
    #             'Label': 'Left consolidation',
    #             'Volume': np.sum((masks_lobe < 3) * (masks_ggo_consolidation == 1)) * voxel_volume
    #         },
    #         'RightConsolidation': {
    #             'Label': 'Right consolidation',
    #             'Volume': np.sum((masks_lobe >= 3) * (masks_ggo_consolidation == 1)) * voxel_volume
    #         }
    #     }

    # Template
    # filename = args.out if args.out != '' else args.exam
    filename = f"{patient_id}_{unique_id}"

    print('variables', variables)
    template = Template(filename=f"{templates_folder}/index.html")
    empty_snapshot = k3d.plot().get_snapshot(9)

    with open(f"{reports_folder}/{patient_id}_{unique_id}.json", "w") as f:
        f.write(json.dumps(variables, indent=4))




    # PDF

    doc_filename = f"{reports_folder}/{patient_id}_{unique_id}.docx"
    screenshots = {}

    for plot, name, _ in plots:
        screenshots = {**screenshots, **get_screenshots(plot, name.capitalize(), 8080)}

    # dump
    # for k, v in screenshots.items():
    #     with open(k + ".png", "wb") as f:
    #         f.write(v)

    doc = DocxTemplate(f"{templates_folder}/template2.docx")

    # for plot, name, _ in plots:
    #     key = name.capitalize()
    #     doc.pics_to_replace[key + '1'] = screenshots[key + '_z']
    #     doc.pics_to_replace[key + '2'] = screenshots[key + '_x']
    #     doc.pics_to_replace[key + '3'] = screenshots[key + '_y']

    #     if key + '_3d' in screenshots:
    #         doc.pics_to_replace[key + '4'] = screenshots[key + '_3d']

    doc.render(variables)

    try:
        doc.pre_processing()
    except Exception:
        # accept missing images (due to conditional sections)
        pass

    doc.docx.save(doc_filename)
    doc.post_processing(doc_filename)


    # now insert the screenshots..
    # first temporarily save the screenshots to disk
    maybe_mkdir(f"{reports_folder}/temp_{unique_id}")
    for k, v in screenshots.items():
        with open(f"{reports_folder}/temp_{unique_id}/{k}.png", "wb") as f:
            f.write(v)        

    doc = docx.Document(doc_filename)

    for i, model in enumerate(models):
        key = model.capitalize()
        # add picture inside table
        table = doc.tables[2 * (i + 1)]

        cell = table.rows[0].cells[0]
        p = cell.add_paragraph()
        r = p.add_run()
        r.add_picture(f"{reports_folder}/temp_{unique_id}/{key}_z.png", width=docx.shared.Inches(2), height=docx.shared.Inches(2))

        cell = table.rows[0].cells[1]
        p = cell.add_paragraph()
        r = p.add_run()
        r.add_picture(f"{reports_folder}/temp_{unique_id}/{key}_x.png", width=docx.shared.Inches(2), height=docx.shared.Inches(2))

        cell = table.rows[1].cells[0]
        p = cell.add_paragraph()
        r = p.add_run()
        r.add_picture(f"{reports_folder}/temp_{unique_id}/{key}_y.png", width=docx.shared.Inches(2), height=docx.shared.Inches(2))

        if key + '_3d' in screenshots:
            cell = table.rows[1].cells[1]
            p = cell.add_paragraph()
            r = p.add_run()
            r.add_picture(f"{reports_folder}/temp_{unique_id}/{key}_3d.png", width=docx.shared.Inches(2), height=docx.shared.Inches(2))


    doc.save(doc_filename)      
    os.system(f"rm -rf {reports_folder}/temp_{unique_id}")


    try:
        from docx2pdf import convert

        convert(doc_filename, f"{reports_folder}/{patient_id}_{unique_id}.pdf")
    except Exception:
        try:
            # subprocess.call("libreoffice", "--headless --convert-to pdf " + doc_filename)
            os.system(f"libreoffice --headless --convert-to pdf {doc_filename}")
            os.system(f"mv {filename}.pdf {doc_filename.replace('.docx', '.pdf')}")
        except Exception:
            pass





    first = 1
    for i, _, _ in plots:
        for obj in i.objects:
            print(obj.name)
            if obj.name == 'scan':
                if first == 1:
                    first = 0
                    obj.name = 'scan_data'
                else:
                    obj.volume[:, :, :] = 0

    snapshot = empty_snapshot
    FFLATE_JS = re.findall(r'id=\'fflatejs\'>(.+?)</script>', snapshot, re.MULTILINE | re.DOTALL)[0]
    REQUIRE_JS = re.findall(r'id=\'requirejs\'>(.+?)</script>', snapshot, re.MULTILINE | re.DOTALL)[0]
    K3D_SOURCE = re.findall(r'window.k3dCompressed = \'(.+?)\'', snapshot, re.MULTILINE | re.DOTALL)[0]

    variables['FFLATE_JS'] = FFLATE_JS
    variables['REQUIRE_JS'] = REQUIRE_JS
    variables['K3D_SOURCE'] = K3D_SOURCE


    widgets_data = {
        'plots': {}
    }

    for plot, name, type in plots:
        snapshot = plot.get_binary_snapshot(1)
        print(name, len(snapshot) / (1024 ** 2))
        k3d_data = base64.b64encode(snapshot).decode("utf-8")

        widgets_data['plots'][name] = {
            'data': k3d_data,
            'type': type
        }

    variables['json_widgets_data'] = json.dumps(widgets_data)
    variables['json_widgets_data'] = json.dumps(widgets_data)

    with open(f'{templates_folder}/temp_' + filename + '.html', 'wb') as f:
        f.write(template.render_unicode(**variables).encode())

    baked_template = webpage2html.generate(f'{templates_folder}/temp_' + filename + '.html', comment=False,
                                        keep_script=True)

    # cdata fix after webpage2html
    baked_template = baked_template.replace('&gt;', '>').replace('&lt;', '<')

    with open(f"{reports_folder}/{patient_id}_{unique_id}.html", "wb") as f:
        f.write(baked_template.encode())

    # os.remove('./template/temp_' + filename + '.html')
    os.remove(f'{templates_folder}/temp_' + filename + '.html')
    #
    # pdfkit.from_url(filename + '.html', filename + '.pdf', options={
    #     "window-status": "ready_to_print",
    #     "print-media-type": None
    # })


if __name__ == '__main__':

    import firebase_admin
    from firebase_admin import firestore

    patient_id = '3003'
    unique_id = 'b5a48fd02823493b8acc8794c51a93ec'
    models = ['lungs', 'lobes', 'infiltration', 'ggo_consolidation', 'fibrosis_4labels', 'fibrosis_7labels', 'haa']
    niftis_folder = '/lungs/niftis'
    uploads_folder = '/lungs/uploads/'
    outputs_folder = '/lungs/outputs/'
    niftis_folder = '/lungs/niftis/'
    reports_folder = '/lungs/reports/'
    templates_folder = '/lungs/templates/'
    # niftis_folder = '../../data/niftis/'
    # reports_folder = '../../data/reports/'
    # templates_folder = 'templates/'
    # uploads_folder = '../../data/uploads/'
    # outputs_folder = '../../data/outputs/'

    config = yaml.load(open('config.yaml'), Loader=yaml.Loader)
    # config = yaml.load(open('../config.yaml'), Loader=yaml.Loader)

    firebase_cred = firebase_admin.credentials.Certificate('chestomx-firebase.json')
    firebase_admin.initialize_app(firebase_cred)
    db = firestore.client()
    col = db.collection('predictionRecords')

    patient_info = col.document(unique_id).get().to_dict()
    ctss = patient_info['ctss']

  
    make_report(
        patient_id=patient_id,
        unique_id=unique_id,
        models=models,
        model_labels=config['labels'],
        ignore=config['ignore'],
        niftis_folder=niftis_folder,
        uploads_folder=uploads_folder,
        outputs_folder=outputs_folder,
        templates_folder=templates_folder,
        reports_folder=reports_folder,
        patient_info=patient_info,
        ctss=ctss

    )
